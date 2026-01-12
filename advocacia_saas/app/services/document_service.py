"""
Serviço para extração de texto de documentos (PDF, DOCX).
Usado para análise de documentos jurídicos com IA.
"""

from typing import Dict, Optional, Tuple

from werkzeug.datastructures import FileStorage


def extract_text_from_pdf(file: FileStorage) -> Tuple[str, Dict]:
    """
    Extrai texto de um arquivo PDF.

    Args:
        file: Arquivo PDF (FileStorage do Flask)

    Returns:
        Tuple[str, Dict]: (texto extraído, metadados)
    """
    try:
        from PyPDF2 import PdfReader  # type: ignore

        # Ler o PDF
        pdf_reader = PdfReader(file)

        text_parts = []
        for page_num, page in enumerate(pdf_reader.pages, 1):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"--- Página {page_num} ---\n{page_text}")

        full_text = "\n\n".join(text_parts)

        metadata = {
            "pages": len(pdf_reader.pages),
            "characters": len(full_text),
            "format": "pdf",
        }

        return full_text, metadata

    except Exception as e:
        raise Exception(f"Erro ao extrair texto do PDF: {str(e)}")


def extract_text_from_docx(file: FileStorage) -> Tuple[str, Dict]:
    """
    Extrai texto de um arquivo DOCX.

    Args:
        file: Arquivo DOCX (FileStorage do Flask)

    Returns:
        Tuple[str, Dict]: (texto extraído, metadados)
    """
    try:
        from docx import Document  # type: ignore

        # Ler o DOCX
        doc = Document(file)

        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Também extrair texto de tabelas
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)

        full_text = "\n\n".join(text_parts)

        metadata = {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "characters": len(full_text),
            "format": "docx",
        }

        return full_text, metadata

    except Exception as e:
        raise Exception(f"Erro ao extrair texto do DOCX: {str(e)}")


def extract_text_from_txt(file: FileStorage) -> Tuple[str, Dict]:
    """
    Extrai texto de um arquivo TXT.

    Args:
        file: Arquivo TXT (FileStorage do Flask)

    Returns:
        Tuple[str, Dict]: (texto extraído, metadados)
    """
    try:
        # Tentar diferentes encodings
        content = file.read()

        for encoding in ["utf-8", "latin-1", "cp1252"]:
            try:
                text = content.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        else:
            text = content.decode("utf-8", errors="ignore")

        metadata = {
            "characters": len(text),
            "lines": text.count("\n") + 1,
            "format": "txt",
        }

        return text, metadata

    except Exception as e:
        raise Exception(f"Erro ao extrair texto do TXT: {str(e)}")


def extract_document_text(file: FileStorage) -> Tuple[str, Dict]:
    """
    Extrai texto de um documento baseado na extensão.

    Args:
        file: Arquivo (FileStorage do Flask)

    Returns:
        Tuple[str, Dict]: (texto extraído, metadados)

    Raises:
        ValueError: Se o formato não for suportado
    """
    filename = file.filename.lower()

    if filename.endswith(".pdf"):
        return extract_text_from_pdf(file)
    elif filename.endswith(".docx"):
        return extract_text_from_docx(file)
    elif filename.endswith(".doc"):
        raise ValueError(
            "Formato .doc não suportado. Por favor, converta para .docx ou .pdf"
        )
    elif filename.endswith(".txt"):
        return extract_text_from_txt(file)
    else:
        raise ValueError(
            f"Formato não suportado: {filename}. Use PDF, DOCX ou TXT."
        )


def get_supported_formats() -> list:
    """Retorna lista de formatos suportados"""
    return ["pdf", "docx", "txt"]


def validate_document_file(file: FileStorage) -> Tuple[bool, Optional[str]]:
    """
    Valida se o arquivo é um documento válido para análise.

    Args:
        file: Arquivo para validar

    Returns:
        Tuple[bool, str]: (é válido, mensagem de erro se inválido)
    """
    if not file or not file.filename:
        return False, "Nenhum arquivo enviado"

    filename = file.filename.lower()
    supported = get_supported_formats()

    extension = filename.rsplit(".", 1)[-1] if "." in filename else ""

    if extension not in supported:
        return False, f"Formato não suportado. Use: {', '.join(supported)}"

    # Verificar tamanho (máximo 10MB)
    file.seek(0, 2)  # Vai para o final
    size = file.tell()
    file.seek(0)  # Volta para o início

    max_size = 10 * 1024 * 1024  # 10MB
    if size > max_size:
        return False, "Arquivo muito grande. Máximo: 10MB"

    return True, None
